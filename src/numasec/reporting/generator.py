"""
NumaSec - Report Generator

Orchestrates report generation with template rendering and export.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any

from numasec.compliance import (
    CVSSResult,
    PTESPhase,
)


# ══════════════════════════════════════════════════════════════════════════════
# Enums & Configuration
# ══════════════════════════════════════════════════════════════════════════════


class ReportFormat(str, Enum):
    """Report output format."""

    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "markdown"
    HTML = "html"
    JSON = "json"


class ReportTemplate(str, Enum):
    """Report template type."""

    PTES = "ptes"
    EXECUTIVE = "executive"
    TECHNICAL = "technical"
    COMPLIANCE = "compliance"


class Classification(str, Enum):
    """Document classification level."""

    PUBLIC = "Public"
    INTERNAL = "Internal"
    CONFIDENTIAL = "Confidential"
    RESTRICTED = "Restricted"


class OverallRisk(str, Enum):
    """Overall risk rating."""

    CRITICAL = "Critical"
    HIGH = "High"
    MEDIUM = "Medium"
    LOW = "Low"
    MINIMAL = "Minimal"


# ══════════════════════════════════════════════════════════════════════════════
# Data Classes
# ══════════════════════════════════════════════════════════════════════════════


@dataclass
class BrandingConfig:
    """Branding configuration for reports."""

    company_name: str = "NumaSec"
    logo_path: str | None = None
    primary_color: str = "#1a1a2e"
    accent_color: str = "#e94560"
    font_family: str = "Noto Sans"


@dataclass
class VersionEntry:
    """Version history entry."""

    version: str
    date: str
    changes: str
    author: str


@dataclass
class DistributionEntry:
    """Distribution list entry."""

    name: str
    role: str
    organization: str


@dataclass
class EvidenceItem:
    """Evidence for a finding."""

    title: str
    description: str
    code: str | None = None
    screenshot_path: str | None = None


@dataclass
class FindingData:
    """Finding data for report."""

    id: str
    title: str
    severity: str
    cvss_score: float
    cvss_vector: str
    cwe_id: str
    description: str
    affected_assets: list[str]
    evidence: list[EvidenceItem]
    impact: str
    remediation: str
    references: list[str]


@dataclass
class RemediationItem:
    """Remediation roadmap item."""

    priority: int
    finding: str
    effort: str
    timeline: str


@dataclass
class ImplementationGuidance:
    """Implementation guidance section."""

    title: str
    description: str


@dataclass
class ToolUsed:
    """Tool used in assessment."""

    name: str
    purpose: str


@dataclass
class ScopeDetail:
    """Scope detail entry."""

    type: str
    target: str
    status: str


@dataclass
class ReportConfig:
    """Report configuration."""

    # Document metadata
    title: str
    client_name: str
    project_name: str
    version: str = "1.0"
    classification: Classification = Classification.CONFIDENTIAL
    author: str = "NumaSec"
    reviewer: str = ""
    report_date: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d"))

    # Engagement dates
    start_date: str = ""
    end_date: str = ""

    # Template settings
    template: ReportTemplate = ReportTemplate.PTES
    format: ReportFormat = ReportFormat.PDF
    branding: BrandingConfig = field(default_factory=BrandingConfig)

    # Content sections
    executive_summary: str = ""
    scope_items: list[str] = field(default_factory=list)
    testing_approach: str = ""

    # Findings
    findings: list[FindingData] = field(default_factory=list)

    # Metadata
    version_history: list[VersionEntry] = field(default_factory=list)
    distribution_list: list[DistributionEntry] = field(default_factory=list)
    tools_used: list[ToolUsed] = field(default_factory=list)

    # Remediation
    key_recommendations: list[dict[str, str]] = field(default_factory=list)
    remediation_roadmap: list[RemediationItem] = field(default_factory=list)
    implementation_guidance: list[ImplementationGuidance] = field(default_factory=list)

    # Scope details
    scope_details: list[ScopeDetail] = field(default_factory=list)


@dataclass
class ReportResult:
    """Result of report generation."""

    success: bool
    output_path: Path | None
    format: ReportFormat
    errors: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    generation_time: float = 0.0


# ══════════════════════════════════════════════════════════════════════════════
# Report Generator
# ══════════════════════════════════════════════════════════════════════════════


class ReportGenerator:
    """
    Orchestrates report generation.

    Supports:
    - Multiple templates (PTES, Executive, Technical)
    - Multiple formats (PDF, DOCX, Markdown, HTML, JSON)
    - Template variable injection
    - Branding customization
    """

    TEMPLATES_DIR = Path(__file__).parent / "templates"

    def __init__(self, config: ReportConfig) -> None:
        """Initialize generator with config."""
        self.config = config
        self._template_content: str | None = None

    def generate(self, output_path: Path) -> ReportResult:
        """
        Generate report.

        Args:
            output_path: Output file path

        Returns:
            ReportResult with status and path
        """
        import time

        start_time = time.time()
        errors: list[str] = []
        warnings: list[str] = []

        try:
            # Generate based on format
            if self.config.format == ReportFormat.PDF:
                result = self._generate_pdf(output_path)
            elif self.config.format == ReportFormat.DOCX:
                result = self._generate_docx(output_path)
            elif self.config.format == ReportFormat.MARKDOWN:
                result = self._generate_markdown(output_path)
            elif self.config.format == ReportFormat.HTML:
                result = self._generate_html(output_path)
            elif self.config.format == ReportFormat.JSON:
                result = self._generate_json(output_path)
            else:
                errors.append(f"Unsupported format: {self.config.format}")
                result = None

            elapsed = time.time() - start_time

            if result:
                return ReportResult(
                    success=True,
                    output_path=result,
                    format=self.config.format,
                    warnings=warnings,
                    generation_time=elapsed,
                )
            else:
                return ReportResult(
                    success=False,
                    output_path=None,
                    format=self.config.format,
                    errors=errors,
                    generation_time=elapsed,
                )

        except Exception as e:
            elapsed = time.time() - start_time
            return ReportResult(
                success=False,
                output_path=None,
                format=self.config.format,
                errors=[str(e)],
                generation_time=elapsed,
            )

    def _get_template_path(self) -> Path:
        """Get template file path."""
        template_name = self.config.template.value
        return self.TEMPLATES_DIR / template_name / "report.typ"

    def _prepare_context(self) -> dict[str, Any]:
        """Prepare template context from config."""
        # Count findings by severity
        severity_counts = {
            "critical": 0,
            "high": 0,
            "medium": 0,
            "low": 0,
            "informational": 0,
        }

        for finding in self.config.findings:
            severity_lower = finding.severity.lower()
            if severity_lower in severity_counts:
                severity_counts[severity_lower] += 1
            elif severity_lower == "info":
                severity_counts["informational"] += 1

        # Calculate overall risk
        overall_risk = self._calculate_overall_risk(severity_counts)

        # Calculate duration
        duration = self._calculate_duration()

        # Build context
        return {
            "title": self.config.title,
            "client_name": self.config.client_name,
            "project_name": self.config.project_name,
            "version": self.config.version,
            "classification": self.config.classification.value,
            "author": self.config.author,
            "reviewer": self.config.reviewer,
            "report_date": self.config.report_date,
            "start_date": self.config.start_date,
            "end_date": self.config.end_date,
            "duration": duration,
            "status": "Final",
            # Findings counts
            "critical_count": severity_counts["critical"],
            "high_count": severity_counts["high"],
            "medium_count": severity_counts["medium"],
            "low_count": severity_counts["low"],
            "info_count": severity_counts["informational"],
            # Risk
            "overall_risk": overall_risk,
            "overall_risk_color": self._get_risk_color(overall_risk),
            # Content
            "executive_summary": self.config.executive_summary,
            "scope_items": self.config.scope_items,
            "testing_approach": self.config.testing_approach,
            # Lists
            "findings": [self._finding_to_dict(f) for f in self.config.findings],
            "version_history": [
                {"version": v.version, "date": v.date, "changes": v.changes, "author": v.author}
                for v in self.config.version_history
            ],
            "distribution_list": [
                {"name": d.name, "role": d.role, "organization": d.organization}
                for d in self.config.distribution_list
            ],
            "tools_used": [
                {"name": t.name, "purpose": t.purpose}
                for t in self.config.tools_used
            ],
            "key_recommendations": self.config.key_recommendations,
            "remediation_roadmap": [
                {"priority": r.priority, "finding": r.finding, "effort": r.effort, "timeline": r.timeline}
                for r in self.config.remediation_roadmap
            ],
            "implementation_guidance": [
                {"title": g.title, "description": g.description}
                for g in self.config.implementation_guidance
            ],
            "scope_details": [
                {"type": s.type, "target": s.target, "status": s.status}
                for s in self.config.scope_details
            ],
        }

    def _finding_to_dict(self, finding: FindingData) -> dict[str, Any]:
        """Convert finding to dict for template."""
        return {
            "id": finding.id,
            "title": finding.title,
            "severity": finding.severity,
            "cvss_score": finding.cvss_score,
            "cvss_vector": finding.cvss_vector,
            "cwe_id": finding.cwe_id,
            "description": finding.description,
            "affected_assets": finding.affected_assets,
            "evidence": [
                {
                    "title": e.title,
                    "description": e.description,
                    "code": e.content,
                }
                for e in finding.evidence
            ],
            "impact": finding.impact,
            "remediation": finding.remediation,
            "references": finding.references,
        }

    def _calculate_overall_risk(self, counts: dict[str, int]) -> str:
        """Calculate overall risk based on findings."""
        if counts["critical"] > 0:
            return "Critical"
        elif counts["high"] >= 3:
            return "Critical"
        elif counts["high"] > 0:
            return "High"
        elif counts["medium"] >= 5:
            return "High"
        elif counts["medium"] > 0:
            return "Medium"
        elif counts["low"] > 0:
            return "Low"
        else:
            return "Minimal"

    def _get_risk_color(self, risk: str) -> str:
        """Get Typst color for risk level."""
        colors = {
            "Critical": "rgb(\"#8B0000\")",
            "High": "rgb(\"#FF4500\")",
            "Medium": "rgb(\"#FFA500\")",
            "Low": "rgb(\"#32CD32\")",
            "Minimal": "rgb(\"#1E90FF\")",
        }
        return colors.get(risk, "rgb(\"#808080\")")

    def _calculate_duration(self) -> int:
        """Calculate engagement duration in days."""
        try:
            start = datetime.strptime(self.config.start_date, "%Y-%m-%d")
            end = datetime.strptime(self.config.end_date, "%Y-%m-%d")
            return (end - start).days + 1
        except (ValueError, TypeError):
            return 0

    def _render_template(self, template_content: str, context: dict[str, Any]) -> str:
        """
        Render Typst template with Handlebars-like syntax.

        Supports:
        - {{variable}} - Simple substitution
        - {{#each items}} ... {{/each}} - Iteration
        - {{#if condition}} ... {{/if}} - Conditionals
        """
        content = template_content

        # Handle {{#each}} blocks
        each_pattern = r'\{\{#each\s+(\w+)\}\}(.*?)\{\{/each\}\}'
        
        def replace_each(match: re.Match) -> str:
            var_name = match.group(1)
            block_content = match.group(2)
            items = context.get(var_name, [])
            
            result = []
            for idx, item in enumerate(items):
                rendered = block_content
                
                # Replace {{@index}}
                rendered = rendered.replace("{{@index}}", str(idx + 1))
                
                # Replace {{this}} for simple lists
                if isinstance(item, str):
                    rendered = rendered.replace("{{this}}", item)
                elif isinstance(item, dict):
                    # Replace {{key}} with item values
                    for key, value in item.items():
                        rendered = rendered.replace(f"{{{{{key}}}}}", str(value) if value else "")
                
                result.append(rendered)
            
            return "".join(result)

        content = re.sub(each_pattern, replace_each, content, flags=re.DOTALL)

        # Handle {{#if}} blocks
        if_pattern = r'\{\{#if\s+(\w+)\}\}(.*?)\{\{/if\}\}'
        
        def replace_if(match: re.Match) -> str:
            var_name = match.group(1)
            block_content = match.group(2)
            value = context.get(var_name)
            if value:
                return block_content
            return ""

        content = re.sub(if_pattern, replace_if, content, flags=re.DOTALL)

        # Simple variable substitution
        for key, value in context.items():
            if isinstance(value, (str, int, float)):
                content = content.replace(f"{{{{{key}}}}}", str(value))

        return content

    def _generate_pdf(self, output_path: Path) -> Path | None:
        """Generate PDF using Typst."""
        import subprocess
        import tempfile

        # Load template
        template_path = self._get_template_path()
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        template_content = template_path.read_text()

        # Prepare context and render
        context = self._prepare_context()
        rendered = self._render_template(template_content, context)

        # Write to temp file
        with tempfile.NamedTemporaryFile(mode="w", suffix=".typ", delete=False) as f:
            f.write(rendered)
            temp_typ = Path(f.name)

        try:
            # Run Typst
            result = subprocess.run(
                ["typst", "compile", str(temp_typ), str(output_path)],
                capture_output=True,
                text=True,
                timeout=60,
            )

            if result.returncode != 0:
                raise RuntimeError(f"Typst error: {result.stderr}")

            return output_path

        finally:
            temp_typ.unlink(missing_ok=True)

    def _generate_docx(self, output_path: Path) -> Path | None:
        """Generate DOCX using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = Document()
        context = self._prepare_context()

        # Title
        title = doc.add_heading(context["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph(f"{context['client_name']} - {context['project_name']}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add sections
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(context["executive_summary"])

        doc.add_heading("Findings Summary", level=2)
        table = doc.add_table(rows=2, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Critical"
        hdr[1].text = "High"
        hdr[2].text = "Medium"
        hdr[3].text = "Low"
        hdr[4].text = "Info"
        row = table.rows[1].cells
        row[0].text = str(context["critical_count"])
        row[1].text = str(context["high_count"])
        row[2].text = str(context["medium_count"])
        row[3].text = str(context["low_count"])
        row[4].text = str(context["info_count"])

        # Findings
        doc.add_heading("Detailed Findings", level=1)
        for finding in context["findings"]:
            doc.add_heading(f"Finding: {finding['title']}", level=2)
            doc.add_paragraph(f"Severity: {finding['severity']} | CVSS: {finding['cvss_score']}")
            doc.add_paragraph(finding["description"])
            doc.add_heading("Remediation", level=3)
            doc.add_paragraph(finding["remediation"])

        doc.save(str(output_path))
        return output_path

    def _generate_markdown(self, output_path: Path) -> Path | None:
        """Generate Markdown report."""
        context = self._prepare_context()
        
        lines = [
            f"# {context['title']}",
            "",
            f"**Client:** {context['client_name']}",
            f"**Project:** {context['project_name']}",
            f"**Date:** {context['report_date']}",
            f"**Classification:** {context['classification']}",
            "",
            "---",
            "",
            "## Executive Summary",
            "",
            context["executive_summary"],
            "",
            "### Findings Summary",
            "",
            "| Critical | High | Medium | Low | Info |",
            "|----------|------|--------|-----|------|",
            f"| {context['critical_count']} | {context['high_count']} | {context['medium_count']} | {context['low_count']} | {context['info_count']} |",
            "",
            f"**Overall Risk:** {context['overall_risk']}",
            "",
            "---",
            "",
            "## Detailed Findings",
            "",
        ]

        for i, finding in enumerate(context["findings"], 1):
            lines.extend([
                f"### Finding {i}: {finding['title']}",
                "",
                f"**Severity:** {finding['severity']} | **CVSS:** {finding['cvss_score']} | **{finding['cwe_id']}**",
                "",
                "#### Description",
                "",
                finding["description"],
                "",
                "#### Affected Assets",
                "",
            ])
            for asset in finding["affected_assets"]:
                lines.append(f"- `{asset}`")
            
            lines.extend([
                "",
                "#### Impact",
                "",
                finding["impact"],
                "",
                "#### Remediation",
                "",
                finding["remediation"],
                "",
                "---",
                "",
            ])

        output_path.write_text("\n".join(lines))
        return output_path

    def _generate_html(self, output_path: Path) -> Path | None:
        """Generate HTML report."""
        context = self._prepare_context()
        
        # Basic HTML template
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{context['title']}</title>
    <style>
        body {{ font-family: 'Segoe UI', sans-serif; margin: 40px; background: #f5f5f5; }}
        .container {{ max-width: 1000px; margin: 0 auto; background: white; padding: 40px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        h1 {{ color: #1a1a2e; border-bottom: 3px solid #e94560; padding-bottom: 10px; }}
        h2 {{ color: #16213e; margin-top: 30px; }}
        .meta {{ color: #666; margin-bottom: 20px; }}
        .severity {{ padding: 3px 8px; border-radius: 3px; color: white; font-weight: bold; font-size: 12px; }}
        .critical {{ background: #8B0000; }}
        .high {{ background: #FF4500; }}
        .medium {{ background: #FFA500; color: black; }}
        .low {{ background: #32CD32; color: black; }}
        .finding {{ border: 1px solid #ddd; padding: 20px; margin: 20px 0; border-radius: 5px; }}
        .finding h3 {{ margin-top: 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 10px; text-align: center; }}
        th {{ background: #1a1a2e; color: white; }}
        code {{ background: #f0f0f0; padding: 2px 5px; border-radius: 3px; }}
        .summary-table td:nth-child(1) {{ background: #8B0000; color: white; }}
        .summary-table td:nth-child(2) {{ background: #FF4500; color: white; }}
        .summary-table td:nth-child(3) {{ background: #FFA500; }}
        .summary-table td:nth-child(4) {{ background: #32CD32; }}
        .summary-table td:nth-child(5) {{ background: #1E90FF; color: white; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{context['title']}</h1>
        <p class="meta">
            <strong>Client:</strong> {context['client_name']} |
            <strong>Project:</strong> {context['project_name']} |
            <strong>Date:</strong> {context['report_date']} |
            <strong>Classification:</strong> {context['classification']}
        </p>
        
        <h2>Executive Summary</h2>
        <p>{context['executive_summary']}</p>
        
        <h2>Findings Summary</h2>
        <table class="summary-table">
            <tr>
                <th>Critical</th>
                <th>High</th>
                <th>Medium</th>
                <th>Low</th>
                <th>Info</th>
            </tr>
            <tr>
                <td>{context['critical_count']}</td>
                <td>{context['high_count']}</td>
                <td>{context['medium_count']}</td>
                <td>{context['low_count']}</td>
                <td>{context['info_count']}</td>
            </tr>
        </table>
        <p><strong>Overall Risk:</strong> <span class="severity {context['overall_risk'].lower()}">{context['overall_risk']}</span></p>
        
        <h2>Detailed Findings</h2>
"""

        for i, finding in enumerate(context["findings"], 1):
            severity_class = finding['severity'].lower()
            html += f"""
        <div class="finding">
            <h3>Finding {i}: {finding['title']}</h3>
            <p>
                <span class="severity {severity_class}">{finding['severity']}</span>
                CVSS: {finding['cvss_score']} | {finding['cwe_id']}
            </p>
            <h4>Description</h4>
            <p>{finding['description']}</p>
            <h4>Affected Assets</h4>
            <ul>
"""
            for asset in finding["affected_assets"]:
                html += f"                <li><code>{asset}</code></li>\n"
            
            html += f"""            </ul>
            <h4>Impact</h4>
            <p>{finding['impact']}</p>
            <h4>Remediation</h4>
            <p>{finding['remediation']}</p>
        </div>
"""

        html += """
    </div>
</body>
</html>"""

        output_path.write_text(html)
        return output_path

    def _generate_json(self, output_path: Path) -> Path | None:
        """Generate JSON report."""
        context = self._prepare_context()
        output_path.write_text(json.dumps(context, indent=2, default=str))
        return output_path


# ══════════════════════════════════════════════════════════════════════════════
# Helper Functions
# ══════════════════════════════════════════════════════════════════════════════


def create_default_config(
    client_name: str,
    project_name: str,
    findings: list[FindingData],
) -> ReportConfig:
    """Create a default report configuration."""
    return ReportConfig(
        title=f"Penetration Test Report - {project_name}",
        client_name=client_name,
        project_name=project_name,
        findings=findings,
        version="1.0",
        executive_summary=f"This report presents the findings from the penetration test conducted for {client_name}.",
        scope_items=[f"All assets related to {project_name}"],
        testing_approach="Black-box testing approach following PTES methodology.",
        version_history=[
            VersionEntry(
                version="1.0",
                date=datetime.now().strftime("%Y-%m-%d"),
                changes="Initial release",
                author="NumaSec",
            )
        ],
        tools_used=[
            ToolUsed("Nmap", "Network discovery and port scanning"),
            ToolUsed("Nuclei", "Vulnerability scanning"),
            ToolUsed("FFuf", "Web fuzzing"),
            ToolUsed("SQLMap", "SQL injection testing"),
        ],
    )
