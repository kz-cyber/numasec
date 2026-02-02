"""
NumaSec - DOCX Exporter

Generates DOCX reports using python-docx.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any


@dataclass
class DOCXExportResult:
    """Result of DOCX export."""

    success: bool
    output_path: Path | None
    error: str | None = None


class DOCXExporter:
    """
    DOCX exporter using python-docx.

    Creates professional Word documents from report data.
    """

    def __init__(self) -> None:
        """Initialize exporter."""
        self._doc = None

    def _check_docx_available(self) -> bool:
        """Check if python-docx is available."""
        try:
            import docx
            return True
        except ImportError:
            return False

    def export(
        self,
        data: dict[str, Any],
        output_path: Path,
        *,
        template_path: Path | None = None,
    ) -> DOCXExportResult:
        """
        Export data to DOCX.

        Args:
            data: Report data dictionary
            output_path: Output DOCX path
            template_path: Optional template to use

        Returns:
            DOCXExportResult
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            return DOCXExportResult(
                success=False,
                output_path=None,
                error="python-docx not installed. Run: pip install python-docx",
            )

        try:
            # Create document
            if template_path and template_path.exists():
                doc = Document(str(template_path))
            else:
                doc = Document()

            self._doc = doc

            # Setup styles
            self._setup_styles(doc)

            # Add content
            self._add_cover_page(doc, data)
            self._add_executive_summary(doc, data)
            self._add_findings_summary(doc, data)
            self._add_detailed_findings(doc, data)
            self._add_remediation(doc, data)

            # Save
            doc.save(str(output_path))

            return DOCXExportResult(
                success=True,
                output_path=output_path,
            )

        except Exception as e:
            return DOCXExportResult(
                success=False,
                output_path=None,
                error=str(e),
            )

    def _setup_styles(self, doc) -> None:
        """Setup document styles."""
        from docx.shared import Pt, RGBColor
        
        # Try to modify existing styles
        try:
            # Heading 1
            style = doc.styles["Heading 1"]
            style.font.size = Pt(18)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        except KeyError:
            pass

        try:
            # Heading 2
            style = doc.styles["Heading 2"]
            style.font.size = Pt(14)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
        except KeyError:
            pass

    def _add_cover_page(self, doc, data: dict[str, Any]) -> None:
        """Add cover page."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        # Title
        title = doc.add_heading(data.get("title", "Penetration Test Report"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{data.get('client_name', '')} - {data.get('project_name', '')}")
        run.font.size = Pt(14)

        # Date
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get("report_date", datetime.now().strftime("%Y-%m-%d")))
        run.font.size = Pt(12)

        # Classification
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get("classification", "CONFIDENTIAL"))
        run.bold = True

        doc.add_page_break()

    def _add_executive_summary(self, doc, data: dict[str, Any]) -> None:
        """Add executive summary section."""
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(data.get("executive_summary", ""))

    def _add_findings_summary(self, doc, data: dict[str, Any]) -> None:
        """Add findings summary table."""
        doc.add_heading("Findings Summary", level=2)

        # Create table
        table = doc.add_table(rows=2, cols=5)
        table.style = "Table Grid"

        # Header row
        headers = ["Critical", "High", "Medium", "Low", "Info"]
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # Data row
        data_cells = table.rows[1].cells
        data_cells[0].text = str(data.get("critical_count", 0))
        data_cells[1].text = str(data.get("high_count", 0))
        data_cells[2].text = str(data.get("medium_count", 0))
        data_cells[3].text = str(data.get("low_count", 0))
        data_cells[4].text = str(data.get("info_count", 0))

        # Overall risk
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Overall Risk: ").bold = True
        p.add_run(data.get("overall_risk", "Unknown"))

    def _add_detailed_findings(self, doc, data: dict[str, Any]) -> None:
        """Add detailed findings section."""
        doc.add_page_break()
        doc.add_heading("Detailed Findings", level=1)

        findings = data.get("findings", [])
        for i, finding in enumerate(findings, 1):
            self._add_finding(doc, finding, i)

    def _add_finding(self, doc, finding: dict[str, Any], index: int) -> None:
        """Add a single finding."""
        from docx.shared import Pt

        # Finding title
        title = f"Finding {index}: {finding.get('title', 'Unknown')}"
        doc.add_heading(title, level=2)

        # Metadata
        p = doc.add_paragraph()
        p.add_run("Severity: ").bold = True
        p.add_run(finding.get("severity", "Unknown"))
        p.add_run(" | ")
        p.add_run("CVSS: ").bold = True
        p.add_run(str(finding.get("cvss_score", 0.0)))
        p.add_run(" | ")
        p.add_run(finding.get("cwe_id", ""))

        # Description
        doc.add_heading("Description", level=3)
        doc.add_paragraph(finding.get("description", ""))

        # Affected Assets
        affected = finding.get("affected_assets", [])
        if affected:
            doc.add_heading("Affected Assets", level=3)
            for asset in affected:
                doc.add_paragraph(asset, style="List Bullet")

        # Evidence
        evidence_list = finding.get("evidence", [])
        if evidence_list:
            doc.add_heading("Evidence", level=3)
            for evidence in evidence_list:
                p = doc.add_paragraph()
                p.add_run(evidence.get("title", "")).bold = True
                doc.add_paragraph(evidence.get("description", ""))
                if evidence.get("code"):
                    code_para = doc.add_paragraph()
                    run = code_para.add_run(evidence["code"])
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)

        # Impact
        doc.add_heading("Impact", level=3)
        doc.add_paragraph(finding.get("impact", ""))

        # Remediation
        doc.add_heading("Remediation", level=3)
        doc.add_paragraph(finding.get("remediation", ""))

        # References
        refs = finding.get("references", [])
        if refs:
            doc.add_heading("References", level=3)
            for ref in refs:
                doc.add_paragraph(ref, style="List Bullet")

    def _add_remediation(self, doc, data: dict[str, Any]) -> None:
        """Add remediation roadmap."""
        roadmap = data.get("remediation_roadmap", [])
        if not roadmap:
            return

        doc.add_page_break()
        doc.add_heading("Remediation Roadmap", level=1)

        table = doc.add_table(rows=len(roadmap) + 1, cols=4)
        table.style = "Table Grid"

        # Header
        headers = ["Priority", "Finding", "Effort", "Timeline"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        # Data
        for i, item in enumerate(roadmap, 1):
            table.rows[i].cells[0].text = str(item.get("priority", ""))
            table.rows[i].cells[1].text = item.get("finding", "")
            table.rows[i].cells[2].text = item.get("effort", "")
            table.rows[i].cells[3].text = item.get("timeline", "")


# ══════════════════════════════════════════════════════════════════════════════
# Helper Functions
# ══════════════════════════════════════════════════════════════════════════════


def export_to_docx(
    data: dict[str, Any],
    output_path: Path | str,
) -> DOCXExportResult:
    """
    Quick helper to export data to DOCX.

    Args:
        data: Report data dictionary
        output_path: Output DOCX path

    Returns:
        DOCXExportResult
    """
    exporter = DOCXExporter()
    return exporter.export(data, Path(output_path))


def is_docx_available() -> bool:
    """Check if python-docx is available."""
    return DOCXExporter()._check_docx_available()
